





'''                  
	To-do:      run a resume through it and fix things.                      
		At this point, the code is purely generated through Gemini.                       
			(July 24, 2025)   










'''                  





import os
import re
from PyPDF2 import PdfReader
from docx import Document
import spacy
from spacy.matcher import Matcher

# --- Configuration ---
# Path to Tesseract executable (only if you're using pytesseract and it's not in your PATH)
# Example for Windows: TESSERACT_PATH = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
TESSERACT_PATH = None # Set this if Tesseract is not in your system's PATH

if TESSERACT_PATH:
    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
else:
    try:
        import pytesseract
    except ImportError:
        print("Pytesseract not found or Tesseract-OCR not installed. OCR functionality will be limited.")
        pytesseract = None # Disable pytesseract if not available

# Load SpaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("SpaCy model 'en_core_web_sm' not found. Please run 'python -m spacy download en_core_web_sm'")
    nlp = None

# --- Helper Functions for Text Extraction ---

def extract_text_from_pdf(pdf_path):
    """
    Extracts text from a digital PDF file.
    Note: This will not work for scanned PDFs (images).
    """
    text = ""
    try:
        with open(pdf_path, 'rb') as file:
            reader = PdfReader(file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n"
    except Exception as e:
        print(f"Error extracting text from PDF {pdf_path}: {e}")
    return text

def extract_text_from_docx(docx_path):
    """
    Extracts text from a .docx file.
    """
    text = ""
    try:
        document = Document(docx_path)
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        print(f"Error extracting text from DOCX {docx_path}: {e}")
    return text

def extract_text_from_txt(txt_path):
    """
    Extracts text from a .txt file.
    """
    text = ""
    try:
        with open(txt_path, 'r', encoding='utf-8') as file:
            text = file.read()
    except Exception as e:
        print(f"Error extracting text from TXT {txt_path}: {e}")
    return text

def extract_text_from_image_or_scanned_pdf(file_path):
    """
    Extracts text from an image or scanned PDF using Tesseract OCR.
    Requires Tesseract-OCR to be installed and pytesseract to be configured.
    For PDFs, you might need to convert PDF pages to images first (e.g., using Pillow or Poppler).
    This is a simplified example.
    """
    if not pytesseract:
        print("Pytesseract is not available. Cannot perform OCR.")
        return ""

    try:
        # If it's a PDF, we assume it's a scanned PDF for OCR.
        # For better handling, you'd typically convert PDF pages to images first.
        # This example directly uses image_to_string, which might work for single-page PDFs
        # or require additional libraries like `Pillow` for multi-page TIFFs from PDFs.
        from PIL import Image # Pillow is used by pytesseract for image handling
        
        if file_path.lower().endswith('.pdf'):
            # For scanned PDFs, a more robust solution involves converting PDF to images
            # and then running OCR on each image. Poppler can help with this.
            # For simplicity, we'll try a direct approach which might not always work for multi-page PDFs
            # without prior image conversion.
            print(f"Attempting OCR on PDF: {file_path}. For multi-page scanned PDFs, consider converting to images first.")
            # A more robust solution would be:
            # from pdf2image import convert_from_path
            # images = convert_from_path(file_path)
            # text = ""
            # for img in images:
            #     text += pytesseract.image_to_string(img)
            # return text
            # For now, let's just attempt a direct read which might fail for complex PDFs.
            return pytesseract.image_to_string(file_path) # pytesseract can sometimes directly process PDFs (needs Ghostscript)
        else: # Assume it's an image file (PNG, JPG, etc.)
            return pytesseract.image_to_string(Image.open(file_path))
    except Exception as e:
        print(f"Error performing OCR on {file_path}: {e}")
        return ""

# --- Resume Parsing Logic ---

class ResumeParser:
    def __init__(self):
        self.nlp = nlp # SpaCy model
        self.matcher = Matcher(self.nlp.vocab) if self.nlp else None

        if self.matcher:
            # Add patterns for common entities (e.g., email, phone)
            self.add_patterns()

    def add_patterns(self):
        # Email pattern
        email_pattern = [{"TEXT": {"REGEX": r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"}}]
        self.matcher.add("EMAIL", [email_pattern])

        # Phone number pattern (very basic, needs refinement for international formats)
        # Allows for common separators like . - ( ) and space, and optional + at start
        phone_pattern = [{"TEXT": {"REGEX": r"(\+?\d{1,3}[-.\s]?)?(\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}"}}]
        self.matcher.add("PHONE", [phone_pattern])
        
        # LinkedIn URL pattern
        linkedin_pattern = [{"TEXT": {"REGEX": r"linkedin\.com/in/[a-zA-Z0-9_-]+"}}]
        self.matcher.add("LINKEDIN", [linkedin_pattern])

        # GitHub URL pattern
        github_pattern = [{"TEXT": {"REGEX": r"github\.com/[a-zA-Z0-9_-]+"}}]
        self.matcher.add("GITHUB", [github_pattern])

        # Simple pattern for years (e.g., for graduation or experience dates)
        year_pattern = [{"TEXT": {"REGEX": r"\b(19|20)\d{2}\b"}}]
        self.matcher.add("YEAR", [year_pattern])


    def parse_resume(self, file_path):
        """
        Parses a resume file and extracts structured information.
        """
        file_extension = os.path.splitext(file_path)[1].lower()
        raw_text = ""

        if file_extension == '.pdf':
            # Try digital extraction first, then OCR if digital extraction yields little text
            digital_text = extract_text_from_pdf(file_path)
            if len(digital_text.strip()) < 50: # Arbitrary threshold for "little text"
                print(f"Digital PDF extraction yielded little text for {file_path}. Attempting OCR...")
                raw_text = extract_text_from_image_or_scanned_pdf(file_path)
            else:
                raw_text = digital_text
        elif file_extension == '.docx':
            raw_text = extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            raw_text = extract_text_from_txt(file_path)
        elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            raw_text = extract_text_from_image_or_scanned_pdf(file_path)
        else:
            return {"error": f"Unsupported file type: {file_extension}"}

        if not raw_text:
            return {"error": "Could not extract text from the resume."}

        # Clean and pre-process text
        cleaned_text = self._clean_text(raw_text)

        # Use SpaCy for advanced NLP tasks
        doc = self.nlp(cleaned_text) if self.nlp else None
        
        parsed_data = {
            "raw_text": raw_text,
            "cleaned_text": cleaned_text,
            "contact_info": self._extract_contact_info(doc, cleaned_text),
            "education": self._extract_education(doc, cleaned_text),
            "experience": self._extract_experience(doc, cleaned_text),
            "skills": self._extract_skills(doc, cleaned_text),
            "projects": self._extract_projects(doc, cleaned_text),
            "summary": self._extract_summary(doc, cleaned_text),
            # Add more fields as needed
        }
        return parsed_data

    def _clean_text(self, text):
        """
        Basic text cleaning: remove extra whitespace, normalize newlines.
        """
        text = re.sub(r'\s+', ' ', text).strip() # Replace multiple whitespaces with single space
        text = text.replace('\n', ' ') # Replace newlines with spaces for a single line text for easier regex
        return text

    def _extract_contact_info(self, doc, text):
        contact_info = {
            "name": None,
            "email": None,
            "phone": None,
            "linkedin": None,
            "github": None,
            "address": None,
            "portfolio": None
        }

        # Use SpaCy NER for name (PERSON entity)
        if doc:
            for ent in doc.ents:
                if ent.label_ == "PERSON" and contact_info["name"] is None:
                    # Heuristic: Often the first PERSON entity at the top is the name
                    # This is a very basic heuristic and needs refinement.
                    if len(ent.text.split()) >= 2: # At least two words for a name
                        contact_info["name"] = ent.text
                        break # Assume the first good match is the name

        # Use matcher for email, phone, LinkedIn, GitHub
        if self.matcher and doc:
            matches = self.matcher(doc)
            for match_id, start, end in matches:
                span = doc[start:end]
                if self.nlp.vocab.strings[match_id] == "EMAIL":
                    contact_info["email"] = span.text
                elif self.nlp.vocab.strings[match_id] == "PHONE":
                    contact_info["phone"] = span.text
                elif self.nlp.vocab.strings[match_id] == "LINKEDIN":
                    contact_info["linkedin"] = span.text
                elif self.nlp.vocab.strings[match_id] == "GITHUB":
                    contact_info["github"] = span.text

        # Fallback regex for email and phone if SpaCy matcher misses or isn't used
        if not contact_info["email"]:
            email_match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
            if email_match:
                contact_info["email"] = email_match.group(0)

        if not contact_info["phone"]:
            phone_match = re.search(r"(\+?\d{1,3}[-.\s]?)?(\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}", text)
            if phone_match:
                contact_info["phone"] = phone_match.group(0)

        # More complex regex for addresses or using pre-trained NER for GPE (Geo Political Entity)
        # For address, often better to rely on pre-trained GPE or custom rules.
        # Example for a simple city, state, zip:
        address_match = re.search(r"([A-Za-z\s]+, [A-Za-z]{2}\s+\d{5}(-\d{4})?)", text)
        if address_match:
            contact_info["address"] = address_match.group(0)

        return contact_info

    def _extract_education(self, doc, text):
        education_entries = []
        # Keywords to look for
        education_keywords = ["education", "university", "college", "degree", "bachelor", "master", "phd", "associate"]
        
        # Simple regex for degrees and universities
        degree_patterns = [
            r"(B\.?S\.?|M\.?S\.?|Ph\.?D\.?|Bachelor(\'s)?|Master(\'s)?|Associate(\'s)?)\s+in\s+([A-Za-z\s]+)",
            r"([A-Za-z\s]+)\s+Degree\s+in\s+([A-Za-z\s]+)",
            r"(MBA|JD|MD)"
        ]
        
        university_patterns = [
            r"\b(University of [A-Za-z\s]+)\b",
            r"\b([A-Za-z\s]+ (University|College|Institute))\b"
        ]

        # Look for sections related to education
        # This is a very basic approach. For robust extraction, you'd define clear sections
        # and parse within them.
        
        # Example of a very basic sectioning idea (not robust for all resumes)
        sections = re.split(r"(Education|Experience|Skills|Projects)", text, flags=re.IGNORECASE)
        education_section_text = ""
        for i, section in enumerate(sections):
            if section.lower().strip() == "education" and i + 1 < len(sections):
                education_section_text = sections[i+1]
                break
        
        if education_section_text:
            # Find degrees
            for pattern in degree_patterns:
                for match in re.finditer(pattern, education_section_text, re.IGNORECASE):
                    degree_text = match.group(0).strip()
                    # Further parse degree_text for specific fields like type, major
                    education_entries.append({"degree_info": degree_text, "university": None, "year": None})
            
            # Find universities
            for pattern in university_patterns:
                for match in re.finditer(pattern, education_section_text, re.IGNORECASE):
                    university_text = match.group(0).strip()
                    # Try to link to existing degree entry or add new
                    found = False
                    for entry in education_entries:
                        if not entry["university"]:
                            entry["university"] = university_text
                            found = True
                            break
                    if not found:
                        education_entries.append({"degree_info": None, "university": university_text, "year": None})

            # Find years (graduation years)
            if self.matcher and doc:
                matches = self.matcher(doc)
                for match_id, start, end in matches:
                    if self.nlp.vocab.strings[match_id] == "YEAR":
                        year = doc[start:end].text
                        # Simple heuristic: if a year is found near an education entry
                        # This needs sophisticated contextual understanding.
                        for entry in education_entries:
                            if not entry["year"] and year in education_section_text: # Very naive check
                                entry["year"] = year
                                break
        
        # More advanced: using SpaCy's sentence segmentation and custom NER for better results
        if doc:
            for sent in doc.sents:
                sent_text = sent.text.lower()
                if any(keyword in sent_text for keyword in education_keywords):
                    # Here you'd apply more specific patterns or custom NER
                    pass # Placeholder for more complex NLP
        
        return education_entries

    def _extract_experience(self, doc, text):
        experience_entries = []
        # Keywords for experience sections
        experience_keywords = ["experience", "employment history", "work history", "professional experience"]

        # Regex for dates (e.g., "Jan 2020 - Dec 2022", "2020 - Present")
        date_pattern = r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?\.?\s*(\d{4})\s*[-–—]\s*(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?\.?\s*(\d{4}|Present|Current)\b"
        
        # Example of finding potential experience blocks (very basic)
        # Look for lines that contain a title, company, and date
        
        # This is where a more structured parsing approach is needed, e.g., identifying "sections"
        # and then applying rules within those sections.
        
        # Simple extraction of Company and Title using NER
        if doc:
            for ent in doc.ents:
                if ent.label_ == "ORG": # Organization often points to company
                    # Very basic: assume an ORG followed by a GPE (location) or a date is a company
                    # This requires more context.
                    pass # Placeholder for more complex NLP
        
        # For now, let's use a very basic pattern matching for Company, Title, Dates
        # This will be highly unreliable without more structure
        experience_block_pattern = r"([A-Za-z\s,.]+)\n([A-Za-z\s,.]+)\n" + date_pattern + r"\n((?:.|\n)*?)(?=\n\n|\Z)"
        
        # This regex tries to capture: Company \n Title \n Dates \n Description (until two newlines or end of text)
        # This assumes a very specific resume format.
        
        # For a truly lightweight solution *without* robust section parsing, this will be hard.
        # A slightly better approach is to iterate through lines and apply rules.
        
        lines = text.split('\n')
        current_experience = None
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check for dates, often indicative of an experience entry
            date_match = re.search(date_pattern, line, re.IGNORECASE)
            if date_match:
                if current_experience: # Save previous if new one starts
                    experience_entries.append(current_experience)
                current_experience = {
                    "title": None,
                    "company": None,
                    "start_date": date_match.group(0),
                    "end_date": None, # Will need to parse start_date further
                    "description": []
                }
                # Attempt to get company and title from lines preceding the date
                # This is highly heuristic and will fail often
                # A better approach involves relative positioning or heading identification
                
            elif current_experience:
                # Add line to description if within an ongoing experience block
                # This assumes descriptions follow dates and are indented or clearly separate
                current_experience["description"].append(line)
        
        if current_experience: # Add the last one
            experience_entries.append(current_experience)
            
        # Post-process to try and find company/title if not found during initial pass
        # This requires more context, often looking for bolded text or specific heading patterns
        for entry in experience_entries:
            # Example: Try to find company/title in lines directly above the date line
            # This would require revisiting the raw_text or having context.
            pass

        return experience_entries

    def _extract_skills(self, doc, text):
        skills = []
        # Keywords for skills sections
        skill_keywords = ["skills", "technical skills", "proficiencies", "technologies"]

        # Simple approach: look for a "Skills" heading and extract lines below it
        # until another major heading or end of document.
        
        # Regex to find a "Skills" section (assuming it's a prominent heading)
        skill_section_match = re.search(r"(Skills|SKILLS|Technical Skills|PROFICIENCIES|TECHNOLOGIES)\n(.+?)(?=\n\n[A-Z][a-z]+|\Z)", text, re.DOTALL | re.IGNORECASE)
        
        if skill_section_match:
            skill_section_text = skill_section_match.group(2)
            # Split by common delimiters like comma, newline, bullet points
            raw_skills = re.split(r'[,\n•-]', skill_section_text)
            for skill in raw_skills:
                skill = skill.strip()
                if skill and len(skill) > 1 and len(skill.split()) < 5: # Filter out very long phrases
                    skills.append(skill)

        # More advanced: using a pre-defined list of skills and matching them
        # Or using custom NER model trained on skills.
        
        return list(set(skills)) # Remove duplicates

    def _extract_projects(self, doc, text):
        projects = []
        # Similar to experience, this needs good sectioning or clear patterns.
        project_keywords = ["projects", "portfolio", "personal projects"]
        
        # Placeholder
        return projects

    def _extract_summary(self, doc, text):
        summary = None
        # Often at the very beginning. Look for common summary headings.
        summary_keywords = ["summary", "profile", "objective"]
        
        # Very basic: assume the first paragraph if it contains certain keywords or is short.
        first_paragraph_match = re.match(r"^\s*([A-Za-z\s,.]+?)\.", text, re.DOTALL)
        if first_paragraph_match:
            potential_summary = first_paragraph_match.group(1).strip()
            if any(keyword in potential_summary.lower() for keyword in summary_keywords) or len(potential_summary.split()) < 100: # Heuristic for length
                summary = potential_summary

        return summary


# --- Main Execution for Testing ---
if __name__ == "__main__":
    parser = ResumeParser()

    # Create dummy resume files for testing
    dummy_pdf_path = "dummy_resume.pdf"
    dummy_docx_path = "dummy_resume.docx"
    dummy_txt_path = "dummy_resume.txt"
    dummy_scanned_pdf_path = "dummy_scanned_resume.pdf" # This would require an actual scanned PDF or image
    dummy_image_path = "dummy_resume.png" # This would require an actual image

    # Generate dummy PDF (requires reportlab, not included in base setup)
    # For a real test, you'd use a real PDF.
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter

        c = canvas.Canvas(dummy_pdf_path, pagesize=letter)
        c.drawString(100, 750, "John Doe")
        c.drawString(100, 730, "johndoe@example.com | (123) 456-7890 | LinkedIn: linkedin.com/in/johndoe")
        c.drawString(100, 710, "Summary:")
        c.drawString(120, 690, "Experienced software engineer with a passion for building scalable web applications.")
        c.drawString(100, 670, "Education:")
        c.drawString(120, 650, "Master of Science in Computer Science, University of Example, 2022")
        c.drawString(120, 630, "Bachelor of Science in Software Engineering, Tech University, 2020")
        c.drawString(100, 610, "Experience:")
        c.drawString(120, 590, "Senior Software Engineer, Acme Corp, Jan 2022 - Present")
        c.drawString(140, 570, "- Developed and maintained backend services.")
        c.drawString(140, 550, "- Led a team of 3 engineers.")
        c.drawString(120, 530, "Software Engineer, Beta Solutions, Jul 2020 - Dec 2021")
        c.drawString(140, 510, "- Implemented new features.")
        c.drawString(100, 490, "Skills:")
        c.drawString(120, 470, "Python, Java, JavaScript, AWS, Docker, Kubernetes, SQL, NoSQL")
        c.save()
        print(f"Generated dummy PDF: {dummy_pdf_path}")
    except ImportError:
        print("ReportLab not installed. Skipping dummy PDF generation. Please use a real PDF for testing.")

    # Generate dummy DOCX
    try:
        document = Document()
        document.add_paragraph("John Doe")
        document.add_paragraph("johndoe@example.com | (123) 456-7890 | LinkedIn: linkedin.com/in/johndoe")
        document.add_paragraph("Summary:")
        document.add_paragraph("Experienced software engineer with a passion for building scalable web applications.")
        document.add_paragraph("Education:")
        document.add_paragraph("Master of Science in Computer Science, University of Example, 2022")
        document.add_paragraph("Bachelor of Science in Software Engineering, Tech University, 2020")
        document.add_paragraph("Experience:")
        document.add_paragraph("Senior Software Engineer, Acme Corp, Jan 2022 - Present")
        document.add_paragraph("- Developed and maintained backend services.")
        document.add_paragraph("- Led a team of 3 engineers.")
        document.add_paragraph("Software Engineer, Beta Solutions, Jul 2020 - Dec 2021")
        document.add_paragraph("- Implemented new features.")
        document.add_paragraph("Skills:")
        document.add_paragraph("Python, Java, JavaScript, AWS, Docker, Kubernetes, SQL, NoSQL")
        document.save(dummy_docx_path)
        print(f"Generated dummy DOCX: {dummy_docx_path}")
    except Exception as e:
        print(f"Error generating dummy DOCX: {e}. Skipping dummy DOCX generation. Please use a real DOCX for testing.")

    # Generate dummy TXT
    with open(dummy_txt_path, "w") as f:
        f.write("""John Doe
johndoe@example.com | (123) 456-7890
LinkedIn: linkedin.com/in/johndoe

Summary:
Experienced software engineer with a passion for building scalable web applications.

Education:
Master of Science in Computer Science, University of Example, 2022
Bachelor of Science in Software Engineering, Tech University, 2020

Experience:
Senior Software Engineer, Acme Corp, Jan 2022 - Present
- Developed and maintained backend services.
- Led a team of 3 engineers.

Software Engineer, Beta Solutions, Jul 2020 - Dec 2021
- Implemented new features.

Skills:
Python, Java, JavaScript, AWS, Docker, Kubernetes, SQL, NoSQL
""")
    print(f"Generated dummy TXT: {dummy_txt_path}")

    # Test parsing
    print("\n--- Parsing Dummy PDF ---")
    if os.path.exists(dummy_pdf_path):
        parsed_data = parser.parse_resume(dummy_pdf_path)
        import json
        print(json.dumps(parsed_data, indent=2))
    else:
        print("Skipping PDF test as dummy PDF not generated.")

    print("\n--- Parsing Dummy DOCX ---")
    if os.path.exists(dummy_docx_path):
        parsed_data = parser.parse_resume(dummy_docx_path)
        import json
        print(json.dumps(parsed_data, indent=2))
    else:
        print("Skipping DOCX test as dummy DOCX not generated.")

    print("\n--- Parsing Dummy TXT ---")
    if os.path.exists(dummy_txt_path):
        parsed_data = parser.parse_resume(dummy_txt_path)
        import json
        print(json.dumps(parsed_data, indent=2))
    else:
        print("Skipping TXT test as dummy TXT not generated.")

    # Example of how you might handle a real scanned PDF or image
    # For this to work, you need to replace `dummy_scanned_pdf_path`
    # and `dummy_image_path` with actual paths to scanned files.
    # print("\n--- Parsing Dummy Scanned PDF (Requires Tesseract) ---")
    # if os.path.exists(dummy_scanned_pdf_path):
    #     parsed_data = parser.parse_resume(dummy_scanned_pdf_path)
    #     import json
    #     print(json.dumps(parsed_data, indent=2))
    # else:
    #     print(f"Skipping scanned PDF test. Please create a scanned PDF at: {dummy_scanned_pdf_path}")

    # print("\n--- Parsing Dummy Image (Requires Tesseract) ---")
    # if os.path.exists(dummy_image_path):
    #     parsed_data = parser.parse_resume(dummy_image_path)
    #     import json
    #     print(json.dumps(parsed_data, indent=2))
    # else:
    #     print(f"Skipping image test. Please create an image at: {dummy_image_path}")

































































































